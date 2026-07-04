"""
Build a complete Machine Learning assignment submission.

The script intentionally avoids downloaded data and heavy external ML
frameworks. It creates small reproducible datasets with fixed random seeds,
implements the required models with NumPy, draws charts with Pillow, and
exports both DOCX and PDF versions of the final academic report.
"""

from __future__ import annotations

import itertools
import math
import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Sequence, Tuple

import numpy as np
from numpy.lib.stride_tricks import sliding_window_view
from PIL import Image, ImageDraw, ImageFont

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as PdfImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from xml.sax.saxutils import escape


ROOT = Path(__file__).resolve().parent
OUTPUT_DIR = ROOT / "ml_assignment_submission"
FIG_DIR = OUTPUT_DIR / "figures"
DOCX_PATH = OUTPUT_DIR / "ml_data_science_assignment_submission.docx"
PDF_PATH = OUTPUT_DIR / "ml_data_science_assignment_submission.pdf"


# ---------------------------------------------------------------------------
# General utilities
# ---------------------------------------------------------------------------


def ensure_dirs() -> None:
    """Create output folders used by the report builder."""
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    """Load a readable local font, falling back to Pillow's default font."""
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Calibri Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Calibri.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> Tuple[int, int]:
    """Return text width and height for a Pillow drawing context."""
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def draw_centered(
    draw: ImageDraw.ImageDraw,
    xy: Tuple[float, float],
    text: str,
    font: ImageFont.ImageFont,
    fill: Tuple[int, int, int] = (30, 30, 30),
) -> None:
    """Draw text centered on the provided point."""
    w, h = text_size(draw, text, font)
    draw.text((xy[0] - w / 2, xy[1] - h / 2), text, fill=fill, font=font)


def standardize_train_test(
    x_train: np.ndarray, x_test: np.ndarray
) -> Tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    """Standardize features with statistics learned from the training split."""
    mean = x_train.mean(axis=0)
    std = x_train.std(axis=0)
    std[std == 0] = 1.0
    return (x_train - mean) / std, (x_test - mean) / std, mean, std


def stratified_split(
    x: np.ndarray, y: np.ndarray, train_ratio: float = 0.75, seed: int = 42
) -> Tuple[np.ndarray, np.ndarray, np.ndarray, np.ndarray]:
    """Create a deterministic stratified train-test split."""
    rng = np.random.default_rng(seed)
    train_idx: List[int] = []
    test_idx: List[int] = []
    for label in np.unique(y):
        idx = np.where(y == label)[0]
        rng.shuffle(idx)
        cut = int(round(len(idx) * train_ratio))
        train_idx.extend(idx[:cut].tolist())
        test_idx.extend(idx[cut:].tolist())
    rng.shuffle(train_idx)
    rng.shuffle(test_idx)
    return x[train_idx], x[test_idx], y[train_idx], y[test_idx]


def stratified_text_split(
    texts: Sequence[str], labels: Sequence[int], train_ratio: float = 0.75, seed: int = 17
) -> Tuple[List[str], List[str], np.ndarray, np.ndarray]:
    """Stratified split for a list of text samples."""
    rng = np.random.default_rng(seed)
    labels_np = np.asarray(labels)
    train_idx: List[int] = []
    test_idx: List[int] = []
    for label in np.unique(labels_np):
        idx = np.where(labels_np == label)[0]
        rng.shuffle(idx)
        cut = int(round(len(idx) * train_ratio))
        train_idx.extend(idx[:cut].tolist())
        test_idx.extend(idx[cut:].tolist())
    rng.shuffle(train_idx)
    rng.shuffle(test_idx)
    return (
        [texts[i] for i in train_idx],
        [texts[i] for i in test_idx],
        labels_np[train_idx],
        labels_np[test_idx],
    )


def one_hot(y: np.ndarray, class_count: int) -> np.ndarray:
    """Convert integer class labels to one-hot encoded rows."""
    out = np.zeros((len(y), class_count), dtype=float)
    out[np.arange(len(y)), y] = 1.0
    return out


def softmax(logits: np.ndarray) -> np.ndarray:
    """Stable softmax transformation."""
    shifted = logits - logits.max(axis=1, keepdims=True)
    exp_values = np.exp(shifted)
    return exp_values / exp_values.sum(axis=1, keepdims=True)


def sigmoid(z: np.ndarray) -> np.ndarray:
    """Numerically stable sigmoid for binary classifiers."""
    z = np.clip(z, -40, 40)
    return 1.0 / (1.0 + np.exp(-z))


def confusion_matrix(y_true: np.ndarray, y_pred: np.ndarray, class_count: int) -> np.ndarray:
    """Build a confusion matrix with rows as actual classes and columns as predictions."""
    cm = np.zeros((class_count, class_count), dtype=int)
    for actual, pred in zip(y_true, y_pred):
        cm[int(actual), int(pred)] += 1
    return cm


def classification_report_values(
    y_true: np.ndarray, y_pred: np.ndarray, class_names: Sequence[str]
) -> Dict[str, object]:
    """Compute common classification metrics without external libraries."""
    class_count = len(class_names)
    cm = confusion_matrix(y_true, y_pred, class_count)
    rows = []
    f1_values = []
    for i, name in enumerate(class_names):
        tp = cm[i, i]
        fp = cm[:, i].sum() - tp
        fn = cm[i, :].sum() - tp
        precision = tp / (tp + fp) if (tp + fp) else 0.0
        recall = tp / (tp + fn) if (tp + fn) else 0.0
        f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) else 0.0
        f1_values.append(f1)
        rows.append(
            {
                "class": name,
                "precision": precision,
                "recall": recall,
                "f1": f1,
                "support": int(cm[i, :].sum()),
            }
        )
    return {
        "accuracy": float((y_true == y_pred).mean()),
        "macro_f1": float(np.mean(f1_values)),
        "confusion": cm,
        "rows": rows,
    }


def pca_2d(x: np.ndarray) -> np.ndarray:
    """Project feature data to two principal components for visualization."""
    centered = x - x.mean(axis=0)
    _, _, vt = np.linalg.svd(centered, full_matrices=False)
    return centered @ vt[:2].T


def moving_average(values: Sequence[float], window: int = 5) -> List[float]:
    """Smooth a metric curve lightly for cleaner plotted lines."""
    values_arr = np.asarray(values, dtype=float)
    if len(values_arr) < window:
        return values_arr.tolist()
    out = []
    for i in range(len(values_arr)):
        start = max(0, i - window + 1)
        out.append(float(values_arr[start : i + 1].mean()))
    return out


# ---------------------------------------------------------------------------
# Simple plotting with Pillow
# ---------------------------------------------------------------------------


COLORS = [
    (54, 105, 201),
    (218, 91, 67),
    (68, 156, 92),
    (146, 93, 183),
    (210, 151, 52),
]


def make_scatter_plot(
    points: np.ndarray,
    labels: np.ndarray,
    class_names: Sequence[str],
    title: str,
    xlabel: str,
    ylabel: str,
    path: Path,
) -> None:
    """Create a two-dimensional scatter plot as a PNG file."""
    width, height = 1400, 880
    left, right, top, bottom = 110, 300, 95, 115
    plot_w = width - left - right
    plot_h = height - top - bottom
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(30, bold=True)
    label_font = load_font(21, bold=False)
    tick_font = load_font(17, bold=False)
    legend_font = load_font(19, bold=False)

    x = points[:, 0]
    y = points[:, 1]
    x_pad = max(0.3, (x.max() - x.min()) * 0.08)
    y_pad = max(0.3, (y.max() - y.min()) * 0.08)
    x_min, x_max = x.min() - x_pad, x.max() + x_pad
    y_min, y_max = y.min() - y_pad, y.max() + y_pad

    def map_x(value: float) -> float:
        return left + (value - x_min) / (x_max - x_min) * plot_w

    def map_y(value: float) -> float:
        return top + (y_max - value) / (y_max - y_min) * plot_h

    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 214, 220), width=2)
    for i in range(6):
        gx = left + i * plot_w / 5
        gy = top + i * plot_h / 5
        draw.line((gx, top, gx, top + plot_h), fill=(235, 238, 242), width=1)
        draw.line((left, gy, left + plot_w, gy), fill=(235, 238, 242), width=1)
        x_tick = x_min + i * (x_max - x_min) / 5
        y_tick = y_max - i * (y_max - y_min) / 5
        draw_centered(draw, (gx, top + plot_h + 24), f"{x_tick:.1f}", tick_font, (75, 75, 75))
        draw.text((24, gy - 10), f"{y_tick:.1f}", font=tick_font, fill=(75, 75, 75))

    for label in np.unique(labels):
        class_points = points[labels == label]
        color = COLORS[int(label) % len(COLORS)]
        for px, py in class_points:
            cx, cy = map_x(px), map_y(py)
            draw.ellipse((cx - 6, cy - 6, cx + 6, cy + 6), fill=color, outline=(255, 255, 255), width=1)

    draw_centered(draw, (width / 2, 43), title, title_font, (25, 43, 69))
    draw_centered(draw, (left + plot_w / 2, height - 40), xlabel, label_font, (55, 55, 55))
    draw.text((18, top - 38), ylabel, font=label_font, fill=(55, 55, 55))

    legend_x = left + plot_w + 45
    legend_y = top + 25
    draw.text((legend_x, legend_y - 35), "Legend", font=load_font(21, bold=True), fill=(35, 35, 35))
    for i, name in enumerate(class_names):
        y_pos = legend_y + i * 42
        draw.ellipse((legend_x, y_pos, legend_x + 18, y_pos + 18), fill=COLORS[i], outline=(255, 255, 255))
        draw.text((legend_x + 30, y_pos - 3), name, font=legend_font, fill=(45, 45, 45))

    img.save(path)


def make_confusion_plot(
    cm: np.ndarray,
    class_names: Sequence[str],
    title: str,
    path: Path,
) -> None:
    """Draw a confusion matrix heatmap as a PNG file."""
    class_count = len(class_names)
    cell = 145
    left, top = 250, 165
    width = left + class_count * cell + 115
    height = top + class_count * cell + 145
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(28, bold=True)
    header_font = load_font(18, bold=True)
    value_font = load_font(30, bold=True)
    small_font = load_font(16)
    max_value = max(1, int(cm.max()))

    draw_centered(draw, (width / 2, 48), title, title_font, (25, 43, 69))
    draw_centered(draw, (left + class_count * cell / 2, top - 72), "Predicted class", header_font, (60, 60, 60))
    draw.text((58, top - 34), "Actual class", font=header_font, fill=(60, 60, 60))

    for i, name in enumerate(class_names):
        draw_centered(draw, (left + i * cell + cell / 2, top - 25), name, small_font, (60, 60, 60))
        draw.text((70, top + i * cell + cell / 2 - 10), name, font=small_font, fill=(60, 60, 60))

    for r in range(class_count):
        for c in range(class_count):
            intensity = cm[r, c] / max_value
            base = int(245 - 115 * intensity)
            fill = (base, base + 8 if base < 245 else base, 255)
            x0 = left + c * cell
            y0 = top + r * cell
            draw.rectangle((x0, y0, x0 + cell, y0 + cell), fill=fill, outline=(210, 214, 220), width=2)
            draw_centered(draw, (x0 + cell / 2, y0 + cell / 2), str(int(cm[r, c])), value_font, (20, 35, 60))

    img.save(path)


def make_line_chart(
    series: Sequence[Tuple[str, Sequence[float], Tuple[int, int, int]]],
    title: str,
    xlabel: str,
    ylabel: str,
    path: Path,
    y_min: float | None = None,
    y_max: float | None = None,
) -> None:
    """Draw a multi-series line chart as a PNG file."""
    width, height = 1400, 820
    left, right, top, bottom = 105, 280, 90, 110
    plot_w = width - left - right
    plot_h = height - top - bottom
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(30, bold=True)
    label_font = load_font(21)
    tick_font = load_font(16)
    legend_font = load_font(18)

    all_values = np.concatenate([np.asarray(values, dtype=float) for _, values, _ in series])
    if y_min is None:
        y_min = float(all_values.min())
    if y_max is None:
        y_max = float(all_values.max())
    if math.isclose(y_min, y_max):
        y_min -= 0.1
        y_max += 0.1
    pad = (y_max - y_min) * 0.08
    y_min -= pad
    y_max += pad
    max_len = max(len(values) for _, values, _ in series)

    def map_x(index: int) -> float:
        if max_len == 1:
            return left
        return left + index / (max_len - 1) * plot_w

    def map_y(value: float) -> float:
        return top + (y_max - value) / (y_max - y_min) * plot_h

    draw.rectangle((left, top, left + plot_w, top + plot_h), outline=(210, 214, 220), width=2)
    for i in range(6):
        gy = top + i * plot_h / 5
        gx = left + i * plot_w / 5
        draw.line((left, gy, left + plot_w, gy), fill=(235, 238, 242), width=1)
        draw.line((gx, top, gx, top + plot_h), fill=(235, 238, 242), width=1)
        y_tick = y_max - i * (y_max - y_min) / 5
        draw.text((25, gy - 9), f"{y_tick:.2f}", font=tick_font, fill=(75, 75, 75))
        x_tick = int(round(1 + i * (max_len - 1) / 5))
        draw_centered(draw, (gx, top + plot_h + 24), str(x_tick), tick_font, (75, 75, 75))

    for name, values, color in series:
        values_arr = list(values)
        xy = [(map_x(i), map_y(v)) for i, v in enumerate(values_arr)]
        if len(xy) > 1:
            draw.line(xy, fill=color, width=4)
        for i in np.linspace(0, len(xy) - 1, min(10, len(xy)), dtype=int):
            x0, y0p = xy[i]
            draw.ellipse((x0 - 4, y0p - 4, x0 + 4, y0p + 4), fill=color)

    draw_centered(draw, (width / 2, 43), title, title_font, (25, 43, 69))
    draw_centered(draw, (left + plot_w / 2, height - 38), xlabel, label_font, (55, 55, 55))
    draw.text((20, top - 38), ylabel, font=label_font, fill=(55, 55, 55))

    legend_x = left + plot_w + 45
    legend_y = top + 25
    draw.text((legend_x, legend_y - 35), "Legend", font=load_font(21, bold=True), fill=(35, 35, 35))
    for i, (name, _, color) in enumerate(series):
        y_pos = legend_y + i * 42
        draw.line((legend_x, y_pos + 9, legend_x + 28, y_pos + 9), fill=color, width=4)
        draw.ellipse((legend_x + 10, y_pos + 1, legend_x + 18, y_pos + 17), fill=color)
        draw.text((legend_x + 42, y_pos - 3), name, font=legend_font, fill=(45, 45, 45))

    img.save(path)


def make_image_grid(images: np.ndarray, labels: np.ndarray, class_names: Sequence[str], path: Path) -> None:
    """Create a small grid of example images for the CNN dataset."""
    cell = 150
    label_h = 36
    cols = 3
    rows = len(class_names)
    width = 145 + cols * cell + 45
    height = rows * (cell + label_h) + 95
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(28, bold=True)
    label_font = load_font(17, bold=True)
    draw_centered(draw, (width / 2, 38), "Synthetic Shape Image Samples", title_font, (25, 43, 69))
    rng = np.random.default_rng(303)
    for r, name in enumerate(class_names):
        idx = np.where(labels == r)[0]
        chosen = rng.choice(idx, size=cols, replace=False)
        draw.text((18, 85 + r * (cell + label_h) + 50), name, font=label_font, fill=(55, 55, 55))
        for c, image_idx in enumerate(chosen):
            array = np.clip(images[image_idx] * 255, 0, 255).astype(np.uint8)
            tile = Image.fromarray(array, mode="L").resize((cell, cell), resample=Image.Resampling.NEAREST).convert("RGB")
            x0 = 145 + c * cell
            y0 = 80 + r * (cell + label_h)
            img.paste(tile, (x0, y0))
            draw.rectangle((x0, y0, x0 + cell, y0 + cell), outline=(210, 214, 220), width=2)
    img.save(path)


# ---------------------------------------------------------------------------
# Q11: classification and clustering
# ---------------------------------------------------------------------------


def make_flower_dataset(seed: int = 11) -> Tuple[np.ndarray, np.ndarray, List[str]]:
    """Generate a small flower-measurement dataset with three natural groups."""
    rng = np.random.default_rng(seed)
    class_names = ["Type A", "Type B", "Type C"]
    means = np.array(
        [
            [5.0, 3.45, 1.45, 0.25],
            [5.9, 2.85, 4.25, 1.30],
            [6.65, 3.05, 5.55, 2.05],
        ]
    )
    scales = np.array(
        [
            [0.30, 0.25, 0.20, 0.08],
            [0.38, 0.27, 0.34, 0.16],
            [0.42, 0.30, 0.40, 0.18],
        ]
    )
    data = []
    labels = []
    for label, (mean, scale) in enumerate(zip(means, scales)):
        samples = rng.normal(mean, scale, size=(60, 4))
        samples = np.clip(samples, a_min=0.05, a_max=None)
        data.append(samples)
        labels.extend([label] * len(samples))
    return np.vstack(data), np.asarray(labels), class_names


def train_softmax_classifier(
    x_train: np.ndarray,
    y_train: np.ndarray,
    x_test: np.ndarray,
    y_test: np.ndarray,
    class_count: int,
    epochs: int = 700,
    lr: float = 0.10,
    l2: float = 0.001,
) -> Dict[str, object]:
    """Train a multiclass softmax classifier with gradient descent."""
    rng = np.random.default_rng(21)
    w = rng.normal(0, 0.03, size=(x_train.shape[1], class_count))
    b = np.zeros(class_count)
    y_oh = one_hot(y_train, class_count)
    train_loss: List[float] = []
    train_acc: List[float] = []
    test_acc: List[float] = []

    for _ in range(epochs):
        logits = x_train @ w + b
        probs = softmax(logits)
        loss = -np.sum(y_oh * np.log(probs + 1e-9)) / len(x_train) + 0.5 * l2 * np.sum(w * w)
        grad_logits = (probs - y_oh) / len(x_train)
        grad_w = x_train.T @ grad_logits + l2 * w
        grad_b = grad_logits.sum(axis=0)
        w -= lr * grad_w
        b -= lr * grad_b

        train_pred = np.argmax(x_train @ w + b, axis=1)
        test_pred = np.argmax(x_test @ w + b, axis=1)
        train_loss.append(float(loss))
        train_acc.append(float((train_pred == y_train).mean()))
        test_acc.append(float((test_pred == y_test).mean()))

    y_pred = np.argmax(x_test @ w + b, axis=1)
    return {
        "weights": w,
        "bias": b,
        "y_pred": y_pred,
        "train_loss": train_loss,
        "train_acc": train_acc,
        "test_acc": test_acc,
    }


def kmeans(x: np.ndarray, k: int = 3, seed: int = 8, max_iter: int = 80) -> Dict[str, object]:
    """Cluster feature data with a simple k-means implementation."""
    rng = np.random.default_rng(seed)
    centers = x[rng.choice(len(x), size=k, replace=False)].copy()
    labels = np.zeros(len(x), dtype=int)
    for _ in range(max_iter):
        distances = ((x[:, None, :] - centers[None, :, :]) ** 2).sum(axis=2)
        new_labels = distances.argmin(axis=1)
        new_centers = centers.copy()
        for cluster_id in range(k):
            cluster_points = x[new_labels == cluster_id]
            if len(cluster_points):
                new_centers[cluster_id] = cluster_points.mean(axis=0)
        if np.array_equal(new_labels, labels):
            labels = new_labels
            centers = new_centers
            break
        labels = new_labels
        centers = new_centers
    inertia = float(((x - centers[labels]) ** 2).sum())
    return {"labels": labels, "centers": centers, "inertia": inertia}


def silhouette_score(x: np.ndarray, labels: np.ndarray) -> float:
    """Compute the mean silhouette coefficient for k-means clusters."""
    scores = []
    unique_labels = np.unique(labels)
    for i, point in enumerate(x):
        same = x[labels == labels[i]]
        if len(same) <= 1:
            a = 0.0
        else:
            a = float(np.linalg.norm(same - point, axis=1).sum() / (len(same) - 1))
        b_values = []
        for other_label in unique_labels:
            if other_label == labels[i]:
                continue
            other = x[labels == other_label]
            b_values.append(float(np.linalg.norm(other - point, axis=1).mean()))
        b = min(b_values) if b_values else 0.0
        denom = max(a, b)
        scores.append((b - a) / denom if denom else 0.0)
    return float(np.mean(scores))


def best_cluster_alignment(y_true: np.ndarray, cluster_labels: np.ndarray, class_count: int) -> Tuple[np.ndarray, float]:
    """Map anonymous cluster ids to class ids using the best permutation."""
    best_pred = cluster_labels.copy()
    best_acc = -1.0
    for perm in itertools.permutations(range(class_count)):
        mapped = np.asarray([perm[label] for label in cluster_labels])
        acc = float((mapped == y_true).mean())
        if acc > best_acc:
            best_acc = acc
            best_pred = mapped
    return best_pred, best_acc


def run_q11() -> Dict[str, object]:
    """Execute classification and clustering experiments for Q11."""
    x, y, class_names = make_flower_dataset()
    x_train, x_test, y_train, y_test = stratified_split(x, y, train_ratio=0.75, seed=12)
    x_train_std, x_test_std, mean, std = standardize_train_test(x_train, x_test)
    x_std_all = (x - mean) / std

    classifier = train_softmax_classifier(x_train_std, y_train, x_test_std, y_test, len(class_names))
    class_report = classification_report_values(y_test, classifier["y_pred"], class_names)

    cluster_result = kmeans(x_std_all, k=3, seed=19)
    aligned_clusters, cluster_acc = best_cluster_alignment(y, cluster_result["labels"], len(class_names))
    cluster_cm = confusion_matrix(y, aligned_clusters, len(class_names))
    sil = silhouette_score(x_std_all, cluster_result["labels"])

    pca_points = pca_2d(x_std_all)
    true_scatter = FIG_DIR / "q11_pca_true_classes.png"
    cluster_scatter = FIG_DIR / "q11_pca_kmeans_clusters.png"
    class_cm_path = FIG_DIR / "q11_classification_confusion.png"
    cluster_cm_path = FIG_DIR / "q11_cluster_confusion.png"
    make_scatter_plot(pca_points, y, class_names, "Q11 PCA View: Actual Flower Classes", "Principal Component 1", "Principal Component 2", true_scatter)
    make_scatter_plot(
        pca_points,
        cluster_result["labels"],
        ["Cluster 1", "Cluster 2", "Cluster 3"],
        "Q11 PCA View: K-means Cluster Assignments",
        "Principal Component 1",
        "Principal Component 2",
        cluster_scatter,
    )
    make_confusion_plot(class_report["confusion"], class_names, "Q11 Softmax Classifier Confusion Matrix", class_cm_path)
    make_confusion_plot(cluster_cm, class_names, "Q11 K-means Cluster Alignment Matrix", cluster_cm_path)

    return {
        "dataset_shape": x.shape,
        "class_names": class_names,
        "classification": class_report,
        "classifier_history": classifier,
        "cluster_accuracy": cluster_acc,
        "cluster_silhouette": sil,
        "cluster_inertia": cluster_result["inertia"],
        "cluster_confusion": cluster_cm,
        "figures": {
            "true_scatter": true_scatter,
            "cluster_scatter": cluster_scatter,
            "class_cm": class_cm_path,
            "cluster_cm": cluster_cm_path,
        },
    }


# ---------------------------------------------------------------------------
# Q12: ANN and sentiment classification
# ---------------------------------------------------------------------------


def make_student_success_dataset(seed: int = 32) -> Tuple[np.ndarray, np.ndarray, List[str]]:
    """Generate a binary student-success dataset with interpretable features."""
    rng = np.random.default_rng(seed)
    n = 420
    study_hours = np.clip(rng.normal(4.8, 1.7, n), 0.4, 10.0)
    attendance = np.clip(rng.normal(0.76, 0.12, n), 0.35, 1.0)
    prior_score = np.clip(rng.normal(65, 12, n), 25, 98)
    assignment_score = np.clip(rng.normal(70, 11, n), 20, 100)
    sleep_hours = np.clip(rng.normal(6.7, 0.9, n), 3.5, 9.5)
    x = np.column_stack([study_hours, attendance, prior_score, assignment_score, sleep_hours])
    latent = (
        0.55 * study_hours
        + 3.0 * attendance
        + 0.045 * prior_score
        + 0.038 * assignment_score
        + 0.23 * sleep_hours
        + rng.normal(0, 0.40, n)
    )
    # A median split keeps the synthetic target balanced, so accuracy and F1
    # both reflect performance on the two classes.
    y = (latent > np.median(latent)).astype(int)
    feature_names = ["Study hours", "Attendance", "Prior score", "Assignment score", "Sleep hours"]
    return x, y, feature_names


def train_binary_mlp(
    x_train: np.ndarray,
    y_train: np.ndarray,
    x_test: np.ndarray,
    y_test: np.ndarray,
    hidden: int = 10,
    epochs: int = 750,
    lr: float = 0.045,
    seed: int = 51,
) -> Dict[str, object]:
    """Train a one-hidden-layer ANN for binary classification."""
    rng = np.random.default_rng(seed)
    w1 = rng.normal(0, math.sqrt(2 / x_train.shape[1]), size=(x_train.shape[1], hidden))
    b1 = np.zeros(hidden)
    w2 = rng.normal(0, math.sqrt(2 / hidden), size=(hidden, 1))
    b2 = np.zeros(1)
    y_col = y_train.reshape(-1, 1).astype(float)
    train_loss: List[float] = []
    train_acc: List[float] = []
    test_acc: List[float] = []

    for _ in range(epochs):
        z1 = x_train @ w1 + b1
        a1 = np.maximum(0, z1)
        pred = sigmoid(a1 @ w2 + b2)
        loss = -np.mean(y_col * np.log(pred + 1e-9) + (1 - y_col) * np.log(1 - pred + 1e-9))

        # Back-propagation for binary cross-entropy with sigmoid output.
        dlogit = (pred - y_col) / len(x_train)
        grad_w2 = a1.T @ dlogit
        grad_b2 = dlogit.sum(axis=0)
        da1 = dlogit @ w2.T
        dz1 = da1 * (z1 > 0)
        grad_w1 = x_train.T @ dz1
        grad_b1 = dz1.sum(axis=0)

        w2 -= lr * grad_w2
        b2 -= lr * grad_b2
        w1 -= lr * grad_w1
        b1 -= lr * grad_b1

        train_pred = (sigmoid(np.maximum(0, x_train @ w1 + b1) @ w2 + b2).ravel() >= 0.5).astype(int)
        test_pred = (sigmoid(np.maximum(0, x_test @ w1 + b1) @ w2 + b2).ravel() >= 0.5).astype(int)
        train_loss.append(float(loss))
        train_acc.append(float((train_pred == y_train).mean()))
        test_acc.append(float((test_pred == y_test).mean()))

    y_pred = (sigmoid(np.maximum(0, x_test @ w1 + b1) @ w2 + b2).ravel() >= 0.5).astype(int)
    return {
        "y_pred": y_pred,
        "train_loss": train_loss,
        "train_acc": train_acc,
        "test_acc": test_acc,
    }


STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "but",
    "for",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "the",
    "this",
    "to",
    "was",
    "were",
    "with",
}


POSITIVE_TERMS = {
    "accurate",
    "balanced",
    "bright",
    "careful",
    "clean",
    "cleaner",
    "comfortable",
    "convenient",
    "cool",
    "crisp",
    "dependable",
    "durable",
    "early",
    "easy",
    "encouragement",
    "engaging",
    "enjoyable",
    "excellent",
    "fair",
    "fast",
    "faster",
    "flavor",
    "fresh",
    "friendly",
    "good",
    "happy",
    "helpful",
    "loved",
    "natural",
    "neatly",
    "patience",
    "perfectly",
    "pleasant",
    "polite",
    "practical",
    "quick",
    "quiet",
    "quietly",
    "reassuring",
    "recommend",
    "reliable",
    "respectful",
    "safe",
    "saves",
    "sharp",
    "simple",
    "smooth",
    "soft",
    "solved",
    "strong",
    "sturdy",
    "tasty",
    "thoughtfully",
    "tidy",
    "useful",
    "warm",
    "well",
}


NEGATIVE_TERMS = {
    "avoid",
    "badly",
    "bland",
    "blurry",
    "boring",
    "broken",
    "careless",
    "cluttered",
    "cold",
    "complex",
    "confusing",
    "damaged",
    "dark",
    "delayed",
    "dirty",
    "disliked",
    "disrespectful",
    "dull",
    "failed",
    "fails",
    "frustrating",
    "hard",
    "harsh",
    "heavy",
    "hot",
    "ignored",
    "inaccurate",
    "late",
    "limited",
    "little",
    "loudly",
    "messy",
    "muffled",
    "noisy",
    "poor",
    "poorly",
    "rough",
    "rude",
    "rushed",
    "shaky",
    "shallow",
    "slow",
    "slower",
    "slowly",
    "stale",
    "stressful",
    "struggles",
    "thin",
    "uncomfortable",
    "unclear",
    "unfair",
    "unhappy",
    "unnatural",
    "unpleasant",
    "unreliable",
    "unsafe",
    "unstable",
    "vague",
    "wastes",
    "weak",
}


def make_sentiment_dataset() -> Tuple[List[str], np.ndarray]:
    """Return an original mini dataset of short product and service reviews."""
    positive = [
        "The delivery was quick and the packaging felt careful",
        "I loved the clear sound and comfortable design",
        "The support team solved my issue with patience",
        "This app is smooth reliable and pleasant to use",
        "The meal arrived fresh warm and full of flavor",
        "Installation was simple and the guide was easy",
        "The product quality feels excellent for the price",
        "I am happy with the fast response and helpful staff",
        "The camera captured bright sharp and natural pictures",
        "The course explained difficult ideas in a friendly way",
        "Battery life is strong and the phone stays cool",
        "The fabric is soft durable and neatly stitched",
        "Checkout was quick and the discount worked perfectly",
        "The update made the dashboard cleaner and faster",
        "The headphones are light balanced and enjoyable",
        "The service was polite accurate and reassuring",
        "The room was clean quiet and thoughtfully arranged",
        "I would recommend this tool because it saves time",
        "The feature set is practical and very dependable",
        "The class examples were clear useful and engaging",
        "The charger works well and feels safe",
        "My order arrived early and matched the description",
        "The interface is tidy and easy to understand",
        "The report export was accurate and convenient",
        "The breakfast was tasty fresh and served quickly",
        "The return process was fair simple and respectful",
        "The microphone sounded crisp during the meeting",
        "The laptop runs quietly and handles tasks well",
        "The trainer gave useful feedback and encouragement",
        "The chair is sturdy comfortable and well finished",
    ]
    negative = [
        "The delivery was late and the box was damaged",
        "I disliked the harsh sound and uncomfortable design",
        "The support team ignored my issue for hours",
        "This app is slow unreliable and frustrating to use",
        "The meal arrived cold bland and poorly packed",
        "Installation was confusing and the guide was unclear",
        "The product quality feels poor for the price",
        "I am unhappy with the delayed response and careless staff",
        "The camera captured dark blurry and unnatural pictures",
        "The course rushed difficult ideas without helpful examples",
        "Battery life is weak and the phone becomes hot",
        "The fabric is rough thin and badly stitched",
        "Checkout was slow and the discount failed repeatedly",
        "The update made the dashboard messy and slower",
        "The headphones are heavy dull and unpleasant",
        "The service was rude inaccurate and stressful",
        "The room was dirty noisy and poorly arranged",
        "I would avoid this tool because it wastes time",
        "The feature set is limited and very unstable",
        "The class examples were confusing shallow and boring",
        "The charger fails often and feels unsafe",
        "My order arrived late and did not match the description",
        "The interface is cluttered and hard to understand",
        "The report export was broken and inconvenient",
        "The breakfast was stale cold and served slowly",
        "The return process was unfair complex and disrespectful",
        "The microphone sounded muffled during the meeting",
        "The laptop runs loudly and struggles with tasks",
        "The trainer gave vague feedback and little encouragement",
        "The chair is shaky uncomfortable and poorly finished",
    ]
    texts = positive + negative
    labels = np.asarray([1] * len(positive) + [0] * len(negative), dtype=int)
    return texts, labels


def tokenize(text: str) -> List[str]:
    """Lowercase, remove punctuation, tokenize, and drop common stopwords."""
    tokens = re.findall(r"[a-z]+", text.lower())
    return [token for token in tokens if token not in STOPWORDS and len(token) > 2]


def build_vocab(texts: Sequence[str]) -> Dict[str, int]:
    """Build a vocabulary index from training text only."""
    vocab: Dict[str, int] = {}
    for text in texts:
        for token in tokenize(text):
            if token not in vocab:
                vocab[token] = len(vocab)
    return vocab


def vectorize_texts(texts: Sequence[str], vocab: Dict[str, int]) -> np.ndarray:
    """Convert texts to normalized bag-of-words vectors."""
    x = np.zeros((len(texts), len(vocab)), dtype=float)
    for row, text in enumerate(texts):
        tokens = tokenize(text)
        for token in tokens:
            if token in vocab:
                x[row, vocab[token]] += 1.0
        total = x[row].sum()
        if total > 0:
            x[row] /= total
    return x


def add_sentiment_lexicon_features(texts: Sequence[str], bow: np.ndarray) -> np.ndarray:
    """Add transparent positive/negative word-count features to bag-of-words data."""
    extra_features = []
    for text in texts:
        tokens = tokenize(text)
        positive_hits = sum(token in POSITIVE_TERMS for token in tokens)
        negative_hits = sum(token in NEGATIVE_TERMS for token in tokens)
        extra_features.append([positive_hits, negative_hits, positive_hits - negative_hits])
    return np.hstack([bow, np.asarray(extra_features, dtype=float)])


def run_q12() -> Dict[str, object]:
    """Execute ANN and sentiment-classification experiments for Q12."""
    x, y, feature_names = make_student_success_dataset()
    x_train, x_test, y_train, y_test = stratified_split(x, y, train_ratio=0.76, seed=44)
    x_train_std, x_test_std, _, _ = standardize_train_test(x_train, x_test)
    ann_result = train_binary_mlp(x_train_std, y_train, x_test_std, y_test)
    ann_report = classification_report_values(y_test, ann_result["y_pred"], ["Not selected", "Selected"])

    ann_curve = FIG_DIR / "q12_ann_training_curve.png"
    make_line_chart(
        [
            ("Training loss", moving_average(ann_result["train_loss"], 8), COLORS[1]),
            ("Training accuracy", moving_average(ann_result["train_acc"], 8), COLORS[0]),
            ("Test accuracy", moving_average(ann_result["test_acc"], 8), COLORS[2]),
        ],
        "Q12 ANN Training Progress",
        "Epoch",
        "Metric value",
        ann_curve,
    )
    ann_cm_path = FIG_DIR / "q12_ann_confusion.png"
    make_confusion_plot(ann_report["confusion"], ["Not selected", "Selected"], "Q12 ANN Confusion Matrix", ann_cm_path)

    texts, labels = make_sentiment_dataset()
    train_texts, test_texts, y_train_s, y_test_s = stratified_text_split(texts, labels, train_ratio=0.75, seed=25)
    vocab = build_vocab(train_texts)
    x_train_s = add_sentiment_lexicon_features(train_texts, vectorize_texts(train_texts, vocab))
    x_test_s = add_sentiment_lexicon_features(test_texts, vectorize_texts(test_texts, vocab))
    sentiment_result = train_binary_mlp(
        x_train_s,
        y_train_s,
        x_test_s,
        y_test_s,
        hidden=8,
        epochs=900,
        lr=0.12,
        seed=72,
    )
    sentiment_report = classification_report_values(y_test_s, sentiment_result["y_pred"], ["Negative", "Positive"])
    sentiment_curve = FIG_DIR / "q12_sentiment_training_curve.png"
    sentiment_cm_path = FIG_DIR / "q12_sentiment_confusion.png"
    make_line_chart(
        [
            ("Training loss", moving_average(sentiment_result["train_loss"], 8), COLORS[1]),
            ("Training accuracy", moving_average(sentiment_result["train_acc"], 8), COLORS[0]),
            ("Test accuracy", moving_average(sentiment_result["test_acc"], 8), COLORS[2]),
        ],
        "Q12 Sentiment Classifier Training Progress",
        "Epoch",
        "Metric value",
        sentiment_curve,
    )
    make_confusion_plot(sentiment_report["confusion"], ["Negative", "Positive"], "Q12 Sentiment Confusion Matrix", sentiment_cm_path)

    return {
        "ann_dataset_shape": x.shape,
        "ann_feature_names": feature_names,
        "ann": ann_report,
        "ann_history": ann_result,
        "sentiment_dataset_size": len(texts),
        "sentiment_vocab_size": len(vocab),
        "sentiment": sentiment_report,
        "sentiment_history": sentiment_result,
        "figures": {
            "ann_curve": ann_curve,
            "ann_cm": ann_cm_path,
            "sentiment_curve": sentiment_curve,
            "sentiment_cm": sentiment_cm_path,
        },
    }


# ---------------------------------------------------------------------------
# Q13: CNN for image classification
# ---------------------------------------------------------------------------


def make_shape_image_dataset(seed: int = 91, image_size: int = 24, per_class: int = 180) -> Tuple[np.ndarray, np.ndarray, List[str]]:
    """Generate grayscale images of circles, squares, and triangles."""
    rng = np.random.default_rng(seed)
    class_names = ["Circle", "Square", "Triangle"]
    images = []
    labels = []
    for label, name in enumerate(class_names):
        for _ in range(per_class):
            canvas = Image.new("L", (image_size, image_size), color=0)
            draw = ImageDraw.Draw(canvas)
            size = int(rng.integers(8, 14))
            x0 = int(rng.integers(3, image_size - size - 2))
            y0 = int(rng.integers(3, image_size - size - 2))
            fill = int(rng.integers(190, 256))
            if name == "Circle":
                draw.ellipse((x0, y0, x0 + size, y0 + size), fill=fill)
            elif name == "Square":
                draw.rectangle((x0, y0, x0 + size, y0 + size), fill=fill)
            else:
                draw.polygon(
                    [
                        (x0 + size / 2, y0),
                        (x0, y0 + size),
                        (x0 + size, y0 + size),
                    ],
                    fill=fill,
                )
            array = np.asarray(canvas, dtype=float) / 255.0
            noise = rng.normal(0, 0.08, size=array.shape)
            array = np.clip(array + noise, 0, 1)
            images.append(array)
            labels.append(label)
    return np.asarray(images, dtype=float), np.asarray(labels, dtype=int), class_names


def conv_forward(
    x: np.ndarray, filters: np.ndarray, bias: np.ndarray
) -> Tuple[np.ndarray, Tuple[np.ndarray, Tuple[int, ...], int, int]]:
    """Apply a valid 2D convolution using im2col-style windows."""
    k = filters.shape[1]
    windows = sliding_window_view(x, (k, k), axis=(1, 2))
    batch, out_h, out_w = windows.shape[:3]
    cols = windows.reshape(batch * out_h * out_w, k * k)
    conv = cols @ filters.reshape(filters.shape[0], -1).T + bias
    conv = conv.reshape(batch, out_h, out_w, filters.shape[0]).transpose(0, 3, 1, 2)
    return conv, (cols, x.shape, out_h, out_w)


def conv_backward(
    dout: np.ndarray, cache: Tuple[np.ndarray, Tuple[int, ...], int, int], filter_count: int, kernel_size: int
) -> Tuple[np.ndarray, np.ndarray]:
    """Compute gradients for convolution filters and biases."""
    cols, _, _, _ = cache
    dout_flat = dout.transpose(0, 2, 3, 1).reshape(-1, filter_count)
    grad_filters = dout_flat.T @ cols
    grad_filters = grad_filters.reshape(filter_count, kernel_size, kernel_size)
    grad_bias = dout_flat.sum(axis=0)
    return grad_filters, grad_bias


def maxpool_forward(x: np.ndarray) -> Tuple[np.ndarray, Tuple[np.ndarray, np.ndarray, Tuple[int, ...]]]:
    """Apply non-overlapping 2x2 max pooling."""
    batch, channels, height, width = x.shape
    reshaped = x.reshape(batch, channels, height // 2, 2, width // 2, 2)
    pooled = reshaped.max(axis=(3, 5))
    mask = reshaped == pooled[:, :, :, None, :, None]
    counts = mask.sum(axis=(3, 5), keepdims=True)
    return pooled, (mask, counts, x.shape)


def maxpool_backward(dout: np.ndarray, cache: Tuple[np.ndarray, np.ndarray, Tuple[int, ...]]) -> np.ndarray:
    """Route pooled gradients back to the max locations."""
    mask, counts, original_shape = cache
    distributed = mask * (dout[:, :, :, None, :, None] / counts)
    return distributed.reshape(original_shape)


def cnn_forward(
    x: np.ndarray,
    params: Dict[str, np.ndarray],
) -> Tuple[np.ndarray, Dict[str, object]]:
    """Forward pass for a compact CNN: conv -> ReLU -> pool -> dense -> softmax."""
    conv, conv_cache = conv_forward(x, params["filters"], params["conv_bias"])
    relu_conv = np.maximum(0, conv)
    pooled, pool_cache = maxpool_forward(relu_conv)
    flat = pooled.reshape(len(x), -1)
    hidden_z = flat @ params["w1"] + params["b1"]
    hidden = np.maximum(0, hidden_z)
    logits = hidden @ params["w2"] + params["b2"]
    probs = softmax(logits)
    cache = {
        "conv": conv,
        "conv_cache": conv_cache,
        "pool_cache": pool_cache,
        "pooled_shape": pooled.shape,
        "flat": flat,
        "hidden_z": hidden_z,
        "hidden": hidden,
    }
    return probs, cache


def train_cnn(
    x_train: np.ndarray,
    y_train: np.ndarray,
    x_test: np.ndarray,
    y_test: np.ndarray,
    class_count: int,
    epochs: int = 38,
    batch_size: int = 36,
    lr: float = 0.055,
    seed: int = 125,
) -> Dict[str, object]:
    """Train a small CNN with hand-written NumPy back-propagation."""
    rng = np.random.default_rng(seed)
    filter_count = 5
    kernel_size = 3
    conv_out = x_train.shape[1] - kernel_size + 1
    pooled_side = conv_out // 2
    flat_dim = filter_count * pooled_side * pooled_side
    hidden = 36
    params = {
        "filters": rng.normal(0, 0.12, size=(filter_count, kernel_size, kernel_size)),
        "conv_bias": np.zeros(filter_count),
        "w1": rng.normal(0, math.sqrt(2 / flat_dim), size=(flat_dim, hidden)),
        "b1": np.zeros(hidden),
        "w2": rng.normal(0, math.sqrt(2 / hidden), size=(hidden, class_count)),
        "b2": np.zeros(class_count),
    }
    train_loss: List[float] = []
    train_acc: List[float] = []
    test_acc: List[float] = []
    indices = np.arange(len(x_train))

    for _ in range(epochs):
        rng.shuffle(indices)
        epoch_losses = []
        for start in range(0, len(indices), batch_size):
            batch_idx = indices[start : start + batch_size]
            xb = x_train[batch_idx]
            yb = y_train[batch_idx]
            y_oh = one_hot(yb, class_count)
            probs, cache = cnn_forward(xb, params)
            loss = -np.sum(y_oh * np.log(probs + 1e-9)) / len(xb)
            epoch_losses.append(float(loss))

            # Softmax cross-entropy gradient.
            dlogits = (probs - y_oh) / len(xb)
            grad_w2 = cache["hidden"].T @ dlogits
            grad_b2 = dlogits.sum(axis=0)
            dhidden = dlogits @ params["w2"].T
            dhidden_z = dhidden * (cache["hidden_z"] > 0)
            grad_w1 = cache["flat"].T @ dhidden_z
            grad_b1 = dhidden_z.sum(axis=0)
            dflat = dhidden_z @ params["w1"].T
            dpooled = dflat.reshape(cache["pooled_shape"])
            drelu_conv = maxpool_backward(dpooled, cache["pool_cache"])
            dconv = drelu_conv * (cache["conv"] > 0)
            grad_filters, grad_conv_bias = conv_backward(dconv, cache["conv_cache"], filter_count, kernel_size)

            params["w2"] -= lr * grad_w2
            params["b2"] -= lr * grad_b2
            params["w1"] -= lr * grad_w1
            params["b1"] -= lr * grad_b1
            params["filters"] -= lr * grad_filters
            params["conv_bias"] -= lr * grad_conv_bias

        train_pred = predict_cnn(x_train, params)
        test_pred = predict_cnn(x_test, params)
        train_loss.append(float(np.mean(epoch_losses)))
        train_acc.append(float((train_pred == y_train).mean()))
        test_acc.append(float((test_pred == y_test).mean()))

    y_pred = predict_cnn(x_test, params)
    return {
        "params": params,
        "y_pred": y_pred,
        "train_loss": train_loss,
        "train_acc": train_acc,
        "test_acc": test_acc,
    }


def predict_cnn(x: np.ndarray, params: Dict[str, np.ndarray], batch_size: int = 96) -> np.ndarray:
    """Predict CNN classes in batches."""
    preds = []
    for start in range(0, len(x), batch_size):
        probs, _ = cnn_forward(x[start : start + batch_size], params)
        preds.append(np.argmax(probs, axis=1))
    return np.concatenate(preds)


def run_q13() -> Dict[str, object]:
    """Execute the CNN image-classification experiment for Q13."""
    images, labels, class_names = make_shape_image_dataset()
    x_train, x_test, y_train, y_test = stratified_split(images, labels, train_ratio=0.76, seed=106)
    cnn_result = train_cnn(x_train, y_train, x_test, y_test, len(class_names))
    cnn_report = classification_report_values(y_test, cnn_result["y_pred"], class_names)

    sample_grid = FIG_DIR / "q13_shape_samples.png"
    training_curve = FIG_DIR / "q13_cnn_training_curve.png"
    cnn_cm_path = FIG_DIR / "q13_cnn_confusion.png"
    make_image_grid(images, labels, class_names, sample_grid)
    make_line_chart(
        [
            ("Training loss", moving_average(cnn_result["train_loss"], 3), COLORS[1]),
            ("Training accuracy", moving_average(cnn_result["train_acc"], 3), COLORS[0]),
            ("Test accuracy", moving_average(cnn_result["test_acc"], 3), COLORS[2]),
        ],
        "Q13 CNN Training and Testing Progress",
        "Epoch",
        "Metric value",
        training_curve,
    )
    make_confusion_plot(cnn_report["confusion"], class_names, "Q13 CNN Confusion Matrix", cnn_cm_path)

    return {
        "dataset_shape": images.shape,
        "class_names": class_names,
        "cnn": cnn_report,
        "cnn_history": cnn_result,
        "figures": {
            "samples": sample_grid,
            "curve": training_curve,
            "cm": cnn_cm_path,
        },
    }


# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------


SHORT_ANSWERS = [
    (
        "Q1. Explain the role of Python in Data Science and Machine Learning. Mention any four popular Python libraries used in Machine Learning.",
        [
            "Python is widely used in data science because it is readable, flexible, and supported by a strong ecosystem for data cleaning, statistical analysis, visualization, and model building. In machine learning, Python is used to prepare data, train algorithms, evaluate model performance, automate experiments, and deploy models into applications.",
            "Popular machine learning libraries include NumPy for numerical computation, Pandas for data handling, Scikit-learn for classical machine learning, TensorFlow/Keras for deep learning, PyTorch for neural networks, and Matplotlib for visualization.",
        ],
    ),
    (
        "Q2. Differentiate between supervised learning and unsupervised learning with suitable examples.",
        [
            "Supervised learning uses labelled data, meaning the model learns from examples where the correct output is already known. For example, a model can learn to classify emails as spam or not spam from previously labelled emails.",
            "Unsupervised learning uses unlabelled data and tries to discover hidden patterns. For example, clustering customers into groups based on purchasing behaviour is unsupervised because the group names are not given beforehand.",
        ],
    ),
    (
        "Q3. What is classification in Machine Learning? Name any two classification algorithms.",
        [
            "Classification is a supervised learning task in which a model assigns an input record to one of several predefined categories. Examples include predicting whether a patient is high-risk or low-risk, or identifying whether an image contains a cat, dog, or car.",
            "Two common classification algorithms are Logistic Regression and Decision Trees. Other examples include Support Vector Machines, Naive Bayes, Random Forests, and Neural Networks.",
        ],
    ),
    (
        "Q4. Define clustering. Explain one practical application of clustering.",
        [
            "Clustering is an unsupervised learning technique that groups similar data points together based on patterns in their features. The model does not receive class labels; it forms groups by measuring similarity or distance.",
            "A practical application is customer segmentation. A business can cluster customers by spending level, product preference, and purchase frequency, then design different offers for each segment.",
        ],
    ),
    (
        "Q5. What is overfitting in Machine Learning? How can it be reduced?",
        [
            "Overfitting happens when a model learns the training data too closely, including noise and accidental patterns, and therefore performs poorly on new data. It usually appears as very high training accuracy but much lower testing accuracy.",
            "It can be reduced by using more data, simplifying the model, applying regularization, using cross-validation, stopping training early, pruning decision trees, and applying dropout or data augmentation in neural networks.",
        ],
    ),
    (
        "Q6. Explain the architecture of an Artificial Neural Network (ANN).",
        [
            "An ANN is usually arranged in layers. The input layer receives feature values, one or more hidden layers transform those values through weighted connections and activation functions, and the output layer produces the prediction.",
            "During training, the network compares predicted output with actual output, calculates a loss, and updates weights using back-propagation and gradient descent. The hidden layers allow the ANN to learn nonlinear relationships.",
        ],
    ),
    (
        "Q7. What is the purpose of an activation function in a neural network?",
        [
            "An activation function decides how strongly a neuron should respond to its input. Its most important role is to add nonlinearity, so the neural network can learn complex patterns instead of behaving like a simple linear model.",
            "Common activation functions include ReLU, sigmoid, and tanh. ReLU is widely used in hidden layers, while sigmoid and softmax are often used in output layers for classification tasks.",
        ],
    ),
    (
        "Q8. Differentiate between Deep Learning and traditional Machine Learning.",
        [
            "Traditional machine learning often depends on manually selected features and algorithms such as decision trees, logistic regression, k-means, or support vector machines. It usually works well on structured, moderate-sized datasets.",
            "Deep learning uses neural networks with many layers to learn features automatically from raw or complex data such as images, audio, and text. It often requires more data and computing power but can perform very well on unstructured data.",
        ],
    ),
    (
        "Q9. What is Sentiment Analysis? Mention two real-world applications.",
        [
            "Sentiment analysis is a natural language processing task that identifies the emotional tone or opinion in text, commonly as positive, negative, or neutral. It converts unstructured text into useful opinion-based information.",
            "Two real-world applications are analysing customer reviews to measure product satisfaction and monitoring social media posts to understand public reaction to a brand, policy, event, or campaign.",
        ],
    ),
    (
        "Q10. What is the role of Convolutional Neural Networks (CNNs) in image analysis?",
        [
            "CNNs are designed to analyse image data by using convolution filters that detect local patterns such as edges, corners, textures, and shapes. Pooling layers reduce image size while preserving important visual features.",
            "In image analysis, CNNs are used for image classification, object detection, medical image diagnosis, face recognition, handwriting recognition, and quality inspection in manufacturing.",
        ],
    ),
]


CODE_SNIPPETS = {
    "sklearn_classification": """# Scikit-Learn classification with preprocessing.
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from sklearn.linear_model import LogisticRegression
from sklearn.metrics import accuracy_score, classification_report

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.25, stratify=y, random_state=42
)

classifier = Pipeline([
    ("scaler", StandardScaler()),
    ("model", LogisticRegression(max_iter=1000))
])

classifier.fit(X_train, y_train)
y_pred = classifier.predict(X_test)

print("Accuracy:", accuracy_score(y_test, y_pred))
print(classification_report(y_test, y_pred))""",
    "sklearn_clustering": """# Scikit-Learn k-means clustering and evaluation.
from sklearn.preprocessing import StandardScaler
from sklearn.cluster import KMeans
from sklearn.metrics import silhouette_score

scaler = StandardScaler()
X_scaled = scaler.fit_transform(X)

kmeans = KMeans(n_clusters=3, random_state=42, n_init=10)
cluster_labels = kmeans.fit_predict(X_scaled)

print("Inertia:", kmeans.inertia_)
print("Silhouette score:", silhouette_score(X_scaled, cluster_labels))""",
    "keras_ann": """# Keras/TensorFlow ANN for binary classification.
import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers

ann_model = keras.Sequential([
    layers.Input(shape=(X_train.shape[1],)),
    layers.Dense(10, activation="relu"),
    layers.Dense(1, activation="sigmoid")
])

ann_model.compile(
    optimizer="adam",
    loss="binary_crossentropy",
    metrics=["accuracy"]
)

history = ann_model.fit(
    X_train_scaled, y_train,
    validation_data=(X_test_scaled, y_test),
    epochs=50,
    batch_size=32,
    verbose=1
)""",
    "sklearn_sentiment": """# Scikit-Learn sentiment classifier using text features.
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.metrics import accuracy_score, classification_report

X_train, X_test, y_train, y_test = train_test_split(
    reviews, labels, test_size=0.25, stratify=labels, random_state=42
)

sentiment_model = Pipeline([
    ("tfidf", TfidfVectorizer(stop_words="english")),
    ("classifier", MultinomialNB())
])

sentiment_model.fit(X_train, y_train)
y_pred = sentiment_model.predict(X_test)

print("Accuracy:", accuracy_score(y_test, y_pred))
print(classification_report(y_test, y_pred))""",
    "keras_cnn": """# Keras/TensorFlow CNN for image classification.
from tensorflow import keras
from tensorflow.keras import layers

cnn_model = keras.Sequential([
    layers.Input(shape=(24, 24, 1)),
    layers.Conv2D(16, (3, 3), activation="relu"),
    layers.MaxPooling2D((2, 2)),
    layers.Conv2D(32, (3, 3), activation="relu"),
    layers.MaxPooling2D((2, 2)),
    layers.Flatten(),
    layers.Dense(64, activation="relu"),
    layers.Dense(3, activation="softmax")
])

cnn_model.compile(
    optimizer="adam",
    loss="sparse_categorical_crossentropy",
    metrics=["accuracy"]
)

history = cnn_model.fit(
    X_train_images, y_train,
    validation_data=(X_test_images, y_test),
    epochs=20,
    batch_size=32
)""",
}


def pct(value: float) -> str:
    """Format a decimal as a percentage."""
    return f"{value * 100:.2f}%"


def metric_table_rows(report: Dict[str, object]) -> List[List[str]]:
    """Convert classification-report values to display-table rows."""
    rows = [["Class", "Precision", "Recall", "F1-score", "Support"]]
    for row in report["rows"]:
        rows.append(
            [
                row["class"],
                f"{row['precision']:.3f}",
                f"{row['recall']:.3f}",
                f"{row['f1']:.3f}",
                str(row["support"]),
            ]
        )
    rows.append(["Overall", pct(report["accuracy"]), "Macro F1", f"{report['macro_f1']:.3f}", ""])
    return rows


# ---------------------------------------------------------------------------
# DOCX builder
# ---------------------------------------------------------------------------


def set_cell_shading(cell, fill: str) -> None:
    """Apply background fill to a Word table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    """Set cell margins in twentieths of a point for better table breathing room."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_docx(doc: Document) -> None:
    """Apply the standard_business_brief design preset to the document."""
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_docx_title(doc: Document) -> None:
    """Add the report title block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Machine Learning and Data Science Assignment")
    run.font.name = "Calibri"
    run.font.size = Pt(21)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Original academic submission with Python-based practical work")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("555555")

    doc.add_paragraph()


def add_docx_paragraph(doc: Document, text: str) -> None:
    """Add a normal paragraph to the Word document."""
    doc.add_paragraph(text)


def add_docx_bullets(doc: Document, items: Sequence[str]) -> None:
    """Add bullet items using Word's list style."""
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_docx_table(doc: Document, rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None) -> None:
    """Add a styled table to the Word document."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    header_props = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_props.append(header_marker)
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = table.cell(r, c)
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[c])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9 if r else 9.5)
                    if r == 0:
                        run.font.bold = True
            if r == 0:
                set_cell_shading(cell, "F2F4F7")
    doc.add_paragraph()


def add_docx_figure(doc: Document, path: Path, caption: str, width: float = 6.15) -> None:
    """Insert an inline figure with a caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    # Alt text helps screen readers identify the figure content.
    shape._inline.docPr.set("title", caption.split(".")[0])
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor.from_string("555555")


def add_docx_code(doc: Document, code: str) -> None:
    """Add a compact monospace code block."""
    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def build_docx(q11: Dict[str, object], q12: Dict[str, object], q13: Dict[str, object]) -> None:
    """Create the Word assignment submission."""
    doc = Document()
    style_docx(doc)
    add_docx_title(doc)

    doc.add_heading("Academic Integrity, Assumptions, Datasets, and Libraries", level=1)
    add_docx_paragraph(
        doc,
        "This submission is written in original wording and is intended as a plagiarism-free academic answer. "
        "No external text has been copied. The practical sections use reproducible datasets generated inside Python with fixed random seeds, "
        "so the code, metrics, and figures can be recreated without downloading any private or copyrighted dataset.",
    )
    add_docx_bullets(
        doc,
        [
            "Assumption 1: The practical requirement permits any suitable dataset; therefore, synthetic but realistic datasets are used for classification, clustering, ANN, sentiment analysis, and image classification.",
            "Assumption 2: The goal is to demonstrate the modelling workflow, not to claim production-level performance on real-world data.",
            "Libraries used to generate this report: NumPy for numerical computation, Pillow for figures and generated images, python-docx for the Word file, and ReportLab for the PDF file. The appendix shows standard Scikit-Learn and TensorFlow/Keras code patterns for academic implementation.",
        ],
    )

    doc.add_heading("Section A: Short Answer Questions", level=1)
    for question, answers in SHORT_ANSWERS:
        doc.add_heading(question, level=2)
        for answer in answers:
            add_docx_paragraph(doc, answer)

    doc.add_page_break()
    doc.add_heading("Section B: Descriptive and Practical Questions", level=1)

    doc.add_heading("Q11. Classification and Clustering Using Python", level=2)
    add_docx_paragraph(
        doc,
        "Dataset description: A synthetic flower-measurement dataset was generated with "
        f"{q11['dataset_shape'][0]} rows and {q11['dataset_shape'][1]} numerical features: sepal length, sepal width, petal length, and petal width. "
        "There are three labelled classes. The data resembles a common flower-classification problem, but it is fully generated in this script for originality and reproducibility.",
    )
    add_docx_paragraph(
        doc,
        "Preprocessing steps: The data was split into stratified training and testing sets. Features were standardized using the training mean and standard deviation only. "
        "PCA was then used only for two-dimensional visualization; the models used the standardized feature values.",
    )
    add_docx_paragraph(
        doc,
        "Classification method: A multiclass softmax classifier was trained with gradient descent. Clustering method: k-means was applied with k = 3 on the standardized full dataset.",
    )
    add_docx_table(doc, metric_table_rows(q11["classification"]), widths=[1.45, 1.05, 1.05, 1.05, 0.9])
    add_docx_paragraph(
        doc,
        f"The softmax classifier achieved a test accuracy of {pct(q11['classification']['accuracy'])} and macro F1-score of {q11['classification']['macro_f1']:.3f}. "
        f"K-means achieved an aligned clustering accuracy of {pct(q11['cluster_accuracy'])}, silhouette score of {q11['cluster_silhouette']:.3f}, and inertia of {q11['cluster_inertia']:.2f}.",
    )
    add_docx_figure(doc, q11["figures"]["true_scatter"], "Figure 1. PCA visualization of actual synthetic flower classes.")
    add_docx_figure(doc, q11["figures"]["cluster_scatter"], "Figure 2. PCA visualization of k-means cluster assignments.")
    add_docx_figure(doc, q11["figures"]["class_cm"], "Figure 3. Confusion matrix for the supervised classifier.", width=5.4)
    add_docx_figure(doc, q11["figures"]["cluster_cm"], "Figure 4. Cluster-to-class alignment matrix for k-means.", width=5.4)
    add_docx_paragraph(
        doc,
        "Comparison of outputs: The classifier directly used labelled examples, so its output is a named class for each test record and can be evaluated by accuracy, precision, recall, and F1-score. "
        "K-means did not know the class labels and produced anonymous cluster numbers. After aligning those clusters with the true classes, the result was still useful, but its purpose was pattern discovery rather than direct prediction.",
    )

    doc.add_page_break()
    doc.add_heading("Q12. Artificial Neural Network and Sentiment Classification", level=2)
    add_docx_paragraph(
        doc,
        "ANN dataset description: A synthetic student-success dataset was generated with "
        f"{q12['ann_dataset_shape'][0]} rows and {q12['ann_dataset_shape'][1]} features: "
        + ", ".join(q12["ann_feature_names"])
        + ". The target variable indicates whether the student is selected for an academic support award.",
    )
    add_docx_paragraph(
        doc,
        "ANN architecture: The model used five input neurons, one hidden layer with ten ReLU neurons, and one sigmoid output neuron. "
        "Binary cross-entropy was used as the loss function, and weights were updated by gradient descent.",
    )
    add_docx_table(doc, metric_table_rows(q12["ann"]), widths=[1.55, 1.05, 1.05, 1.05, 0.9])
    add_docx_paragraph(
        doc,
        f"The ANN achieved {pct(q12['ann']['accuracy'])} test accuracy with a macro F1-score of {q12['ann']['macro_f1']:.3f}. "
        "The training curve shows that loss declined while train and test accuracy stabilized, which suggests the model learned a useful decision boundary without extreme overfitting.",
    )
    add_docx_figure(doc, q12["figures"]["ann_curve"], "Figure 5. ANN training loss, training accuracy, and test accuracy.")
    add_docx_figure(doc, q12["figures"]["ann_cm"], "Figure 6. ANN confusion matrix.", width=5.1)

    add_docx_paragraph(
        doc,
        "Sentiment dataset description: A small original review dataset was written for this assignment with "
        f"{q12['sentiment_dataset_size']} short review sentences labelled as positive or negative. The vocabulary built from training data contained "
        f"{q12['sentiment_vocab_size']} terms, and three transparent lexical sentiment features were added: positive-token count, negative-token count, and their difference.",
    )
    add_docx_paragraph(
        doc,
        "NLP preprocessing steps: The text was lowercased, punctuation was removed, sentences were tokenized into words, common stopwords were removed, and each sentence was converted into a normalized bag-of-words vector. "
        "A small positive/negative lexicon was then used to add simple count-based features. The vocabulary was created only from the training split to avoid leakage.",
    )
    add_docx_table(doc, metric_table_rows(q12["sentiment"]), widths=[1.55, 1.05, 1.05, 1.05, 0.9])
    add_docx_paragraph(
        doc,
        f"The sentiment classifier achieved {pct(q12['sentiment']['accuracy'])} test accuracy and macro F1-score of {q12['sentiment']['macro_f1']:.3f}. "
        "Its strengths are simplicity and interpretability. Its limitation is that a small bag-of-words dataset cannot fully capture sarcasm, context, or unseen vocabulary.",
    )
    add_docx_figure(doc, q12["figures"]["sentiment_curve"], "Figure 7. Sentiment classifier training progress.")
    add_docx_figure(doc, q12["figures"]["sentiment_cm"], "Figure 8. Sentiment classifier confusion matrix.", width=5.1)

    doc.add_heading("Q13. CNN for Image Classification", level=2)
    add_docx_paragraph(
        doc,
        "Dataset description: A synthetic image dataset was generated with "
        f"{q13['dataset_shape'][0]} grayscale images of size {q13['dataset_shape'][1]} x {q13['dataset_shape'][2]} pixels. "
        "The three classes were circle, square, and triangle. Random position, size, brightness, and noise were added to make the task less trivial.",
    )
    add_docx_paragraph(
        doc,
        "CNN architecture used: The CNN used one valid 3 x 3 convolution layer with five filters, ReLU activation, 2 x 2 max pooling, one dense hidden layer with 36 ReLU neurons, and a three-neuron softmax output layer.",
    )
    add_docx_paragraph(
        doc,
        "Training and testing process: The image dataset was split in a stratified manner into training and testing sets. The model was trained using mini-batch gradient descent with softmax cross-entropy. "
        "Training loss, training accuracy, and test accuracy were recorded at every epoch.",
    )
    add_docx_table(doc, metric_table_rows(q13["cnn"]), widths=[1.55, 1.05, 1.05, 1.05, 0.9])
    add_docx_paragraph(
        doc,
        f"The CNN achieved {pct(q13['cnn']['accuracy'])} test accuracy and macro F1-score of {q13['cnn']['macro_f1']:.3f}. "
        "The result is strong because the synthetic classes have clear geometric differences. However, the model is limited by the simplicity of the data, small image size, and absence of real-world variation such as shadows, rotations, occlusion, and complex backgrounds.",
    )
    add_docx_figure(doc, q13["figures"]["samples"], "Figure 9. Sample generated images from the CNN dataset.", width=4.9)
    add_docx_figure(doc, q13["figures"]["curve"], "Figure 10. CNN loss and accuracy visualization.")
    add_docx_figure(doc, q13["figures"]["cm"], "Figure 11. CNN confusion matrix.", width=5.2)

    doc.add_heading("Appendix: Commented Python Code Excerpts", level=1)
    add_docx_paragraph(
        doc,
        "The complete reproducible script used to generate the datasets, metrics, charts, DOCX, and PDF is included in the project folder. The appendix below shows standard Scikit-Learn and TensorFlow/Keras code patterns that can be used for a conventional academic implementation.",
    )
    for title, key in [
        ("Scikit-Learn classification excerpt", "sklearn_classification"),
        ("Scikit-Learn clustering excerpt", "sklearn_clustering"),
        ("Keras/TensorFlow ANN excerpt", "keras_ann"),
        ("Scikit-Learn sentiment-analysis excerpt", "sklearn_sentiment"),
        ("Keras/TensorFlow CNN excerpt", "keras_cnn"),
    ]:
        doc.add_heading(title, level=2)
        add_docx_code(doc, CODE_SNIPPETS[key])

    doc.save(DOCX_PATH)


# ---------------------------------------------------------------------------
# PDF builder
# ---------------------------------------------------------------------------


def pdf_styles() -> Dict[str, ParagraphStyle]:
    """Create ReportLab paragraph styles for the PDF output."""
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "Title",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=24,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "Subtitle",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=10.5,
            leading=14,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "H1",
            parent=base["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=15,
            leading=18,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=14,
            spaceAfter=7,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=12,
            leading=15,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=10,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "Body",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=10,
            leading=13.2,
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "caption": ParagraphStyle(
            "Caption",
            parent=base["BodyText"],
            fontName="Helvetica-Oblique",
            fontSize=8.5,
            leading=11,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=8,
        ),
        "code": ParagraphStyle(
            "Code",
            parent=base["Code"],
            fontName="Courier",
            fontSize=7.5,
            leading=9,
            leftIndent=12,
            spaceAfter=8,
        ),
    }


def p(text: str, style: ParagraphStyle) -> Paragraph:
    """Create an escaped PDF paragraph."""
    return Paragraph(escape(text), style)


def bullet(text: str, style: ParagraphStyle) -> Paragraph:
    """Create a simple bullet paragraph."""
    return Paragraph("&bull; " + escape(text), style)


def pdf_metric_table(rows: Sequence[Sequence[str]]) -> Table:
    """Create a styled metrics table for the PDF."""
    table = Table(rows, repeatRows=1, hAlign="LEFT", colWidths=[1.55 * inch, 0.95 * inch, 0.95 * inch, 0.95 * inch, 0.75 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#B8C0CC")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEADING", (0, 0), (-1, -1), 10.5),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    return table


def pdf_figure(path: Path, caption: str, width_in: float = 6.25) -> KeepTogether:
    """Create a PDF image with a caption."""
    img = Image.open(path)
    aspect = img.height / img.width
    pdf_img = PdfImage(str(path), width=width_in * inch, height=width_in * inch * aspect)
    styles = pdf_styles()
    return KeepTogether([pdf_img, p(caption, styles["caption"])])


def build_pdf(q11: Dict[str, object], q12: Dict[str, object], q13: Dict[str, object]) -> None:
    """Create a PDF version of the assignment submission."""
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )
    story: List[object] = []
    story.append(p("Machine Learning and Data Science Assignment", styles["title"]))
    story.append(p("Original academic submission with Python-based practical work", styles["subtitle"]))

    story.append(p("Academic Integrity, Assumptions, Datasets, and Libraries", styles["h1"]))
    story.append(
        p(
            "This submission is written in original wording and is intended as a plagiarism-free academic answer. No external text has been copied. The practical sections use reproducible datasets generated inside Python with fixed random seeds.",
            styles["body"],
        )
    )
    for item in [
        "Assumption 1: Any suitable dataset is permitted; therefore synthetic but realistic datasets are used.",
        "Assumption 2: The goal is to demonstrate the modelling workflow, not production-level performance.",
        "Libraries used to generate the report: NumPy, Pillow, python-docx, and ReportLab. The appendix shows standard Scikit-Learn and TensorFlow/Keras code patterns.",
    ]:
        story.append(bullet(item, styles["body"]))

    story.append(p("Section A: Short Answer Questions", styles["h1"]))
    for question, answers in SHORT_ANSWERS:
        story.append(p(question, styles["h2"]))
        for answer in answers:
            story.append(p(answer, styles["body"]))

    story.append(PageBreak())
    story.append(p("Section B: Descriptive and Practical Questions", styles["h1"]))
    story.append(p("Q11. Classification and Clustering Using Python", styles["h2"]))
    story.append(
        p(
            f"A synthetic flower-measurement dataset was generated with {q11['dataset_shape'][0]} rows and {q11['dataset_shape'][1]} numerical features. The workflow included stratified splitting, standardization, supervised softmax classification, and unsupervised k-means clustering.",
            styles["body"],
        )
    )
    story.append(pdf_metric_table(metric_table_rows(q11["classification"])))
    story.append(Spacer(1, 8))
    story.append(
        p(
            f"The softmax classifier achieved {pct(q11['classification']['accuracy'])} accuracy and macro F1-score {q11['classification']['macro_f1']:.3f}. K-means achieved aligned clustering accuracy {pct(q11['cluster_accuracy'])}, silhouette score {q11['cluster_silhouette']:.3f}, and inertia {q11['cluster_inertia']:.2f}.",
            styles["body"],
        )
    )
    story.append(pdf_figure(q11["figures"]["true_scatter"], "Figure 1. PCA visualization of actual synthetic flower classes."))
    story.append(pdf_figure(q11["figures"]["cluster_scatter"], "Figure 2. PCA visualization of k-means cluster assignments."))
    story.append(pdf_figure(q11["figures"]["class_cm"], "Figure 3. Confusion matrix for the supervised classifier.", 4.6))
    story.append(pdf_figure(q11["figures"]["cluster_cm"], "Figure 4. Cluster-to-class alignment matrix for k-means.", 4.6))
    story.append(
        p(
            "The classifier predicts named classes because it learns from labelled examples. K-means discovers anonymous groups and must be interpreted after training, so it is best treated as exploratory pattern discovery.",
            styles["body"],
        )
    )

    story.append(PageBreak())
    story.append(p("Q12. Artificial Neural Network and Sentiment Classification", styles["h2"]))
    story.append(
        p(
            f"The ANN used a synthetic student-success dataset with {q12['ann_dataset_shape'][0]} rows and {q12['ann_dataset_shape'][1]} features. The architecture was input layer -> 10 ReLU neurons -> sigmoid output.",
            styles["body"],
        )
    )
    story.append(pdf_metric_table(metric_table_rows(q12["ann"])))
    story.append(Spacer(1, 8))
    story.append(
        p(
            f"The ANN achieved {pct(q12['ann']['accuracy'])} test accuracy and macro F1-score {q12['ann']['macro_f1']:.3f}.",
            styles["body"],
        )
    )
    story.append(pdf_figure(q12["figures"]["ann_curve"], "Figure 5. ANN training loss, training accuracy, and test accuracy."))
    story.append(pdf_figure(q12["figures"]["ann_cm"], "Figure 6. ANN confusion matrix.", 4.4))
    story.append(
        p(
            f"The sentiment dataset contained {q12['sentiment_dataset_size']} original review sentences. NLP preprocessing included lowercasing, punctuation removal, tokenization, stopword removal, normalized bag-of-words vectorization, and three lexical sentiment features. The training vocabulary contained {q12['sentiment_vocab_size']} terms.",
            styles["body"],
        )
    )
    story.append(pdf_metric_table(metric_table_rows(q12["sentiment"])))
    story.append(Spacer(1, 8))
    story.append(
        p(
            f"The sentiment classifier achieved {pct(q12['sentiment']['accuracy'])} test accuracy and macro F1-score {q12['sentiment']['macro_f1']:.3f}. Its main limitation is that a small bag-of-words model cannot fully understand context or sarcasm.",
            styles["body"],
        )
    )
    story.append(pdf_figure(q12["figures"]["sentiment_curve"], "Figure 7. Sentiment classifier training progress."))
    story.append(pdf_figure(q12["figures"]["sentiment_cm"], "Figure 8. Sentiment classifier confusion matrix.", 4.4))

    story.append(PageBreak())
    story.append(p("Q13. CNN for Image Classification", styles["h2"]))
    story.append(
        p(
            f"The CNN dataset contained {q13['dataset_shape'][0]} synthetic grayscale images of size {q13['dataset_shape'][1]} x {q13['dataset_shape'][2]} pixels in three classes: circle, square, and triangle.",
            styles["body"],
        )
    )
    story.append(
        p(
            "The CNN architecture was 3 x 3 convolution with five filters, ReLU activation, 2 x 2 max pooling, dense hidden layer with 36 ReLU neurons, and softmax output. Training used mini-batch gradient descent and softmax cross-entropy.",
            styles["body"],
        )
    )
    story.append(pdf_metric_table(metric_table_rows(q13["cnn"])))
    story.append(Spacer(1, 8))
    story.append(
        p(
            f"The CNN achieved {pct(q13['cnn']['accuracy'])} test accuracy and macro F1-score {q13['cnn']['macro_f1']:.3f}. The result is strong because the classes have clear geometric differences, but the dataset is simpler than real image data.",
            styles["body"],
        )
    )
    story.append(pdf_figure(q13["figures"]["samples"], "Figure 9. Sample generated images from the CNN dataset.", 4.4))
    story.append(pdf_figure(q13["figures"]["curve"], "Figure 10. CNN loss and accuracy visualization."))
    story.append(pdf_figure(q13["figures"]["cm"], "Figure 11. CNN confusion matrix.", 4.6))

    story.append(PageBreak())
    story.append(p("Appendix: Commented Python Code Excerpts", styles["h1"]))
    story.append(
        p(
            "The complete reproducible script is included in the project folder. The appendix below shows standard Scikit-Learn and TensorFlow/Keras code patterns for a conventional academic implementation.",
            styles["body"],
        )
    )
    for title, key in [
        ("Scikit-Learn classification excerpt", "sklearn_classification"),
        ("Scikit-Learn clustering excerpt", "sklearn_clustering"),
        ("Keras/TensorFlow ANN excerpt", "keras_ann"),
        ("Scikit-Learn sentiment-analysis excerpt", "sklearn_sentiment"),
        ("Keras/TensorFlow CNN excerpt", "keras_cnn"),
    ]:
        story.append(p(title, styles["h2"]))
        story.append(Preformatted(CODE_SNIPPETS[key], styles["code"]))

    doc.build(story)


def write_summary(q11: Dict[str, object], q12: Dict[str, object], q13: Dict[str, object]) -> None:
    """Write a compact machine-readable summary for quick verification."""
    summary = OUTPUT_DIR / "metrics_summary.txt"
    summary.write_text(
        "\n".join(
            [
                "Generated ML assignment metrics",
                f"Q11 classification accuracy: {pct(q11['classification']['accuracy'])}",
                f"Q11 k-means aligned accuracy: {pct(q11['cluster_accuracy'])}",
                f"Q12 ANN accuracy: {pct(q12['ann']['accuracy'])}",
                f"Q12 sentiment accuracy: {pct(q12['sentiment']['accuracy'])}",
                f"Q13 CNN accuracy: {pct(q13['cnn']['accuracy'])}",
                f"DOCX: {DOCX_PATH}",
                f"PDF: {PDF_PATH}",
            ]
        )
        + "\n",
        encoding="utf-8",
    )


def main() -> None:
    """Run all experiments and build the final deliverables."""
    ensure_dirs()
    q11 = run_q11()
    q12 = run_q12()
    q13 = run_q13()
    build_docx(q11, q12, q13)
    build_pdf(q11, q12, q13)
    write_summary(q11, q12, q13)
    print(f"Created DOCX: {DOCX_PATH}")
    print(f"Created PDF: {PDF_PATH}")
    print(f"Q11 classification accuracy: {pct(q11['classification']['accuracy'])}")
    print(f"Q11 k-means aligned accuracy: {pct(q11['cluster_accuracy'])}")
    print(f"Q12 ANN accuracy: {pct(q12['ann']['accuracy'])}")
    print(f"Q12 sentiment accuracy: {pct(q12['sentiment']['accuracy'])}")
    print(f"Q13 CNN accuracy: {pct(q13['cnn']['accuracy'])}")


if __name__ == "__main__":
    main()
